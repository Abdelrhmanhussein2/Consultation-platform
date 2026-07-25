import io
import re
import uuid
import docx
from helpers.neo4j_db import neo4j_db

class DocxParserService:
    @staticmethod
    def extract_docx_text(file_bytes: bytes) -> str:
        """
        Extracts raw text from both paragraphs and tables of a DOCX file.
        """
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        
        # 1. Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # 2. Extract from tables (useful if laws/articles are formatted in tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
                        
        return '\n'.join(full_text)

    @classmethod
    def resolve_law_name_to_id(cls, law_name: str) -> str:
        """
        Resolves an Arabic law name to its corresponding law_id by searching in Neo4j.
        """
        name_clean = law_name.lower().strip()
        name_clean = re.sub(r'^(قانون|ال)|\s+', '', name_clean)
        
        # Static lookup mappings for common laws
        if "ضريبةالدخل" in name_clean or "ضريبهالدخل" in name_clean:
            return "law_34_2014"
        elif "جمارك" in name_clean:
            return "law_20_1998"
            
        # Dynamic search in Neo4j
        query = """
        MATCH (l:Law)
        WHERE l.title CONTAINS $keyword OR $keyword CONTAINS l.title
        RETURN l.law_id AS law_id
        LIMIT 1
        """
        try:
            keyword = law_name.replace("قانون", "").strip()
            if not keyword:
                return None
            with neo4j_db.get_session() as session:
                res = session.execute_read(lambda tx: tx.run(query, keyword=keyword).single())
                if res:
                    return res["law_id"]
        except Exception:
            pass
        return None

    @classmethod
    def parse_law_docx(cls, file_bytes: bytes, law_id: str, version_name: str, effective_from: str) -> dict:
        """
        Parses a law DOCX file and extracts ArticleVersions, Paragraphs, and Items.
        - Articles are matched using patterns like: المادة 1-, المادة (5), المادة 12:
        - Paragraphs are matched using Arabic alphabetical letters: أ-, ب-, ج)
        - Items are matched using digit numbers: 1-, 2-, 3)
        """
        text = cls.extract_docx_text(file_bytes)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Regex patterns for Arabic legal structure
        article_pattern = re.compile(r'^المادة\s*\(?(\d+)\)?[\-\:\s\.]*')
        paragraph_pattern = re.compile(r'^([أ-ي])[\-\:\s\.\)]+\s*(.*)')
        item_pattern = re.compile(r'^(\d+)[\-\:\s\.\)]+\s*(.*)')
        
        articles = []
        paragraphs = []
        items = []
        
        current_article = None
        current_paragraph = None
        
        for line in lines:
            # 1. Match Article
            art_match = article_pattern.match(line)
            if art_match:
                art_num = int(art_match.group(1))
                art_version_id = f"{law_id}_art_{art_num}_v_{version_name}"
                
                current_article = {
                    "version_id": art_version_id,
                    "text": line,
                    "effective_from": effective_from,
                    "status": "active",
                    "number": art_num,
                    "law_version_id": f"{law_id}_v_{version_name}"
                }
                articles.append(current_article)
                current_paragraph = None
                continue
                
            # 2. Match Paragraph (only if inside an active article)
            if current_article:
                p_match = paragraph_pattern.match(line)
                if p_match:
                    p_letter = p_match.group(1)
                    p_id = f"{current_article['version_id']}_p_{p_letter}"
                    
                    current_paragraph = {
                        "paragraph_id": p_id,
                        "letter": p_letter,
                        "text": line,
                        "article_version_id": current_article["version_id"]
                    }
                    paragraphs.append(current_paragraph)
                    continue
                    
                # 3. Match Item (only if inside an active paragraph)
                if current_paragraph:
                    i_match = item_pattern.match(line)
                    if i_match:
                        i_num = int(i_match.group(1))
                        i_id = f"{current_paragraph['paragraph_id']}_i_{i_num}"
                        
                        item_node = {
                            "item_id": i_id,
                            "number": i_num,
                            "text": line,
                            "paragraph_id": current_paragraph["paragraph_id"]
                        }
                        items.append(item_node)
                        continue
                        
                # 4. Multiline logic: if no structure matches, append text to the current open node
                if current_paragraph:
                    current_paragraph["text"] += "\n" + line
                else:
                    current_article["text"] += "\n" + line
                    
        return {
            "articles": articles,
            "paragraphs": paragraphs,
            "items": items
        }

    @classmethod
    def parse_judgment_docx(
        cls,
        file_bytes: bytes,
        metadata: dict = None,
        default_law_id: str = None,
        default_version_name: str = None
    ) -> dict:
        """
        Parses a judgment DOCX, extracts court metadata using Regex, and finds legal citations.
        """
        text = cls.extract_docx_text(file_bytes)
        
        extracted = {
            "ruling_id": None,
            "case_number": None,
            "ruling_number": None,
            "ruling_year": None,
            "court": None,
            "court_type": "cassation",
            "date": None,
            "outcome": None,
            "subject": None,
            "title": None,
            "full_text": text
        }
        
        # 1. Extract metadata via Regex
        case_match = re.search(r'(?:قضية رقم|رقم القضية|الدعوى رقم)\s*([\d\s]+/[\s\d]+)', text)
        if case_match:
            extracted["case_number"] = case_match.group(1).replace(" ", "")
            
        ruling_match = re.search(r'(?:حكم رقم|قرار رقم)\s*(\d+)(?:\s*/\s*|\s+لسنة\s+)(\d+)', text)
        if ruling_match:
            extracted["ruling_number"] = int(ruling_match.group(1))
            extracted["ruling_year"] = int(ruling_match.group(2))
            
        court_match = re.search(r'محكمة\s+([^\s،؛\.][^،؛\.]+)', text)
        if court_match:
            extracted["court"] = "محكمة " + court_match.group(1).strip()
            
        date_match = re.search(r'(?:تاريخ|تاريخ النطق|بتاريخ)\s*([\d\-\/]{8,10})', text)
        if date_match:
            extracted["date"] = date_match.group(1)
            
        subject_match = re.search(r'(?:الموضوع|موضوع الدعوى)\s*[:\-]\s*([^\n\r]+)', text)
        if subject_match:
            extracted["subject"] = subject_match.group(1).strip()
            
        outcome_match = re.search(r'(?:القرار|النتيجة)\s*[:\-]\s*([^\n\r]+)', text)
        if outcome_match:
            extracted["outcome"] = outcome_match.group(1).strip()

        # Override with user metadata if provided
        if metadata:
            for k, v in metadata.items():
                if v is not None:
                    extracted[k] = v

        # Construct ruling_id if not present
        if not extracted["ruling_id"]:
            if extracted["ruling_number"] and extracted["ruling_year"]:
                extracted["ruling_id"] = f"ruling_{extracted['ruling_number']}_{extracted['ruling_year']}"
            else:
                extracted["ruling_id"] = f"ruling_{uuid.uuid4().hex[:8]}"

        # Deduce court_type
        court_name = (extracted["court"] or "").lower()
        if "تمييز" in court_name:
            extracted["court_type"] = "cassation"
        elif "استئناف" in court_name and "ضريب" in court_name:
            extracted["court_type"] = "appeal_tax"
        elif "بداية" in court_name and "ضريب" in court_name:
            extracted["court_type"] = "tax_first"
            
        if not extracted["title"]:
            extracted["title"] = f"قرار {extracted['court'] or 'المحكمة'} رقم {extracted['ruling_number'] or ''} لسنة {extracted['ruling_year'] or ''}"

        # 2. Extract Citations
        # Pattern matches "المادة 11/و/2 من قانون..." or "المادة (70) من القانون"
        citation_pattern = re.compile(
            r'المادة\s*\(?(\d+)\)?(?:\s*/\s*\(?([أ-ي])\)?)*(?:\s*/\s*\(?(\d+)\)?)*(?:\s+من\s+(قانون\s+[^\s،؛\.][^،؛\.]+))?',
            re.IGNORECASE
        )
        
        citations = []
        for match in citation_pattern.finditer(text):
            art_num = match.group(1)
            p_letter = match.group(2)
            item_num = match.group(3)
            law_name = match.group(4)
            
            # Grab context around citation for Neo4j property storage
            start = max(0, match.start() - 30)
            end = min(len(text), match.end() + 30)
            citation_context = text[start:end].strip().replace("\n", " ")
            
            resolved_law_id = default_law_id
            if law_name:
                resolved_law_id = cls.resolve_law_name_to_id(law_name) or default_law_id
                
            if not resolved_law_id:
                continue
                
            resolved_version_name = default_version_name or "amended"
            
            # Construct specific target ID
            target_id = f"{resolved_law_id}_art_{art_num}_v_{resolved_version_name}"
            target_label = "ArticleVersion"
            
            if p_letter:
                target_id = f"{target_id}_p_{p_letter}"
                target_label = "Paragraph"
                if item_num:
                    target_id = f"{target_id}_i_{item_num}"
                    target_label = "Item"
                    
            citations.append({
                "ruling_id": extracted["ruling_id"],
                "target_id": target_id,
                "target_label": target_label,
                "citation_text": citation_context,
                "paragraph_letter": p_letter,
                "item_number": int(item_num) if item_num else None,
                "law_name": law_name.strip() if law_name else (resolved_law_id.replace("law_", "").replace("_", " "))
            })
            
        return {
            "judgment": extracted,
            "citations": citations
        }
